import streamlit as st
import cohere
import json
import os
import time
import sqlite3
import numpy as np
from io import BytesIO
from docx import Document
from src.pdf_processing import extract_text_from_pdf, split_text
from src.rag_pipeline import (
    create_embeddings,
    build_faiss_index,
    retrieve_chunks,
    generate_answer,
    generate_policy_brief,
)
from src.evaluation import generate_ground_truth, save_ground_truth, evaluate_retrieval
from src.fine_tuning import generate_fine_tune_dataset, start_fine_tuning
from src.storage import init_db, save_chunks_and_embeddings, load_chunks_and_embeddings, load_chunks_for_file, get_file_names

# Set Cohere API key (from secrets, not UI)
COHERE_API_KEY = st.secrets.get("COHERE_API_KEY", "sS5dzpbakUlfsYsIdLMwFUic1MV8JJZbsSpG89Aa")
co = cohere.Client(COHERE_API_KEY)

# Sidebar: File Upload & Model Selection
st.sidebar.title("PV Chatbot Training")
st.sidebar.markdown("Upload multiple PDFs and interact with them using RAG and Cohere.")

DB_PATH = "data/rag.db"
GROUND_TRUTH_PATH = "data/ground_truth.json"
FINE_TUNE_DATASET_PATH = "data/fine_tune_dataset.jsonl"
conn = init_db(DB_PATH)

st.sidebar.subheader("Upload PDFs")
uploaded_files = st.sidebar.file_uploader("Choose PDFs", type="pdf", accept_multiple_files=True)
if uploaded_files:
    existing_file_names = set(get_file_names(conn))
    new_files = []
    for uploaded_file in uploaded_files:
        file_stub = uploaded_file.name.replace('.pdf', '')
        if file_stub in existing_file_names:
            st.sidebar.info(f"{uploaded_file.name} already processed. Using cached content.")
            continue
        new_files.append((uploaded_file, file_stub))

    if new_files:
        for uploaded_file, file_name in new_files:
            with st.spinner(f"Processing {uploaded_file.name}..."):
                text = extract_text_from_pdf(uploaded_file)
                if not text:
                    st.error(f"Failed to extract text from {uploaded_file.name}.")
                    continue
                chunks = split_text(text)
                embeddings = create_embeddings(chunks, co)
                if embeddings.size == 0:
                    st.error(f"Failed to create embeddings for {uploaded_file.name}.")
                    continue
                save_chunks_and_embeddings(conn, file_name, chunks, embeddings)
                st.success(f"{uploaded_file.name} processed ({len(chunks)} chunks).")

    chunks, embeddings, file_names = load_chunks_and_embeddings(conn)
    if chunks:
        unique_files = list(dict.fromkeys(file_names))
        st.session_state['uploaded_files'] = unique_files
        st.session_state['all_chunks'] = chunks
        st.session_state['all_embeddings'] = embeddings
        st.session_state['all_file_names'] = file_names
        st.session_state['index'] = build_faiss_index(embeddings)
        if st.session_state.get('selected_file_choice') not in (["All"] + unique_files):
            st.session_state['selected_file_choice'] = "All"
        st.sidebar.success(
            f"Loaded {len(unique_files)} files from storage. Total chunks: {len(chunks)}"
        )


def extract_policy_template_text(uploaded_template) -> str:
    """Extract readable text from a DOCX policy brief template."""
    try:
        file_bytes = uploaded_template.read()
        if hasattr(uploaded_template, "seek"):
            uploaded_template.seek(0)
        document = Document(BytesIO(file_bytes))
    except Exception as template_error:
        st.sidebar.warning(f"Unable to read template: {template_error}")
        return ""

    lines = []
    for paragraph in document.paragraphs:
        lines.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cell_text = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(filter(None, cell_text)))

    template_text = "\n".join(lines).strip()
    return template_text


st.sidebar.subheader("Policy Brief Template (optional)")
template_uploader = st.sidebar.file_uploader(
    "Upload policy brief template (.docx)",
    type=["docx"],
    key="policy_template_uploader",
)
if template_uploader is not None:
    template_text = extract_policy_template_text(template_uploader)
    st.session_state['policy_template_text'] = template_text
    st.session_state['policy_template_name'] = template_uploader.name
    if template_text:
        st.sidebar.success(f"Loaded template: {template_uploader.name}")
    else:
        st.sidebar.info(
            "Template uploaded but no text detected; the default policy brief outline will be used."
        )

if st.session_state.get('policy_template_text'):
    st.sidebar.caption(
        f"Active template: {st.session_state.get('policy_template_name', 'custom template')}"
    )
    if st.sidebar.button("Remove policy template"):
        st.session_state.pop('policy_template_text', None)
        st.session_state.pop('policy_template_name', None)
        st.sidebar.info("Policy template removed.")


if not st.session_state.get('uploaded_files'):
    chunks, embeddings, file_names = load_chunks_and_embeddings(conn)
    if chunks:
        unique_files = list(dict.fromkeys(file_names))
        st.session_state['uploaded_files'] = unique_files
        st.session_state['all_chunks'] = chunks
        st.session_state['all_embeddings'] = embeddings
        st.session_state['all_file_names'] = file_names
        st.session_state['index'] = build_faiss_index(embeddings)
        if st.session_state.get('selected_file_choice') not in (["All"] + unique_files):
            st.session_state['selected_file_choice'] = "All"
        st.sidebar.info(f"Loaded {len(chunks)} chunks from {len(unique_files)} files.")

# Sidebar: Model Selection
st.sidebar.subheader("Model Selection")
use_finetuned_model = st.sidebar.checkbox("Use Fine-Tuned Model")
fine_tuned_model_id = ""
if use_finetuned_model:
    fine_tuned_model_id = st.sidebar.text_input("Fine-Tuned Model ID:", value=st.session_state.get('fine_tuned_model_id', ""))
    if fine_tuned_model_id:
        st.session_state['fine_tuned_model_id'] = fine_tuned_model_id
else:
    st.session_state['fine_tuned_model_id'] = ""

# Sidebar: Storage Management
st.sidebar.subheader("Storage Management")
# Delete specific PDFs
existing_files = get_file_names(conn)
if existing_files:
    files_to_delete = st.sidebar.multiselect("Delete specific PDFs", options=existing_files)
    if st.sidebar.button("Delete Selected PDFs") and files_to_delete:
        cursor = conn.cursor()
        for fname in files_to_delete:
            cursor.execute('DELETE FROM embeddings WHERE file_name = ?', (fname,))
        conn.commit()
        # Update session state
        chunks, embeddings, file_names = load_chunks_and_embeddings(conn)
        st.session_state['all_chunks'] = chunks
        st.session_state['all_embeddings'] = embeddings
        st.session_state['all_file_names'] = file_names
        st.session_state['index'] = build_faiss_index(embeddings)
        unique_files = list(dict.fromkeys(file_names))
        st.session_state['uploaded_files'] = unique_files
        if st.session_state.get('selected_file_choice') not in (["All"] + unique_files):
            st.session_state['selected_file_choice'] = "All"
        st.sidebar.success(f"Deleted: {', '.join(files_to_delete)}")
# Confirmation for clearing all PDFs
clear_all_confirm = st.sidebar.checkbox("Do you want to clear all PDFs in the database?")
if st.sidebar.button("Clear All PDF Storage") and clear_all_confirm:
    cursor = conn.cursor()
    cursor.execute('DELETE FROM embeddings')
    conn.commit()
    st.session_state['uploaded_files'] = []
    st.session_state['all_chunks'] = []
    st.session_state['all_embeddings'] = np.array([])
    st.session_state['all_file_names'] = []
    st.session_state['index'] = None
    st.session_state['selected_file_choice'] = "All"
    st.sidebar.success("All PDF storage cleared.")

# Main: Workflow Steps
# st.title("PV Chatbot")
# st.markdown("""
# Interact with your PDFs using Retrieval-Augmented Generation and Cohere. We have implemented the following steps below:
# 1. **Upload PDFs** in the sidebar.
# 2. **Generate and Edit Ground Truth** for evaluation.
# 3. **Evaluate Retrieval**.
# 4. **Generate Fine-Tuning Dataset** and **Start Fine-Tuning**.
# 5. **Chat with PDFs** using the best model.
# """)

# Step 3: Chat Interface
st.header("Chat with PVbot")
st.markdown("""
Interact with your policies by following steps below:
1. **Upload Policy PDFs** in the sidebar.
2. **Chat with PVbot** using the best model.
            
""")
# Chatbot UI container
chat_container = st.container()

# Display chat history with avatars
with chat_container:
    if 'chat_history' not in st.session_state or not st.session_state['chat_history']:
        st.chat_message("assistant").write("👋 Hi! Ask me anything about your PDFs.")
    else:
        for msg in st.session_state['chat_history']:
            if msg['role'] == "user":
                st.chat_message("user").write(msg['content'])
            else:
                st.chat_message("assistant").write(msg['content'])

# Chat input at the bottom
# Prepare document selection options
file_options = ["All"]
if st.session_state.get('all_file_names'):
    ordered_unique_files = list(dict.fromkeys(st.session_state['all_file_names']))
    file_options.extend(ordered_unique_files)

# Policy analysis toggle inline with chat input
if 'policy_analysis_mode' not in st.session_state:
    st.session_state['policy_analysis_mode'] = False
if 'selected_file_choice' not in st.session_state:
    st.session_state['selected_file_choice'] = "All"

with st.container():
    current_choice = st.session_state.get('selected_file_choice', 'All')
    if current_choice not in file_options:
        current_choice = 'All'
    selected_index = file_options.index(current_choice)
    selected_file = st.selectbox(
        "Select a policy document to query (or All)",
        options=file_options,
        index=selected_index,
    )
    st.session_state['selected_file_choice'] = selected_file

    policy_mode = st.checkbox(
        "I want an expert policy analysis",
        value=st.session_state.get('policy_analysis_mode', False),
        help="When enabled, responses are formatted as a policy brief using the uploaded template if available.",
    )
    st.session_state['policy_analysis_mode'] = policy_mode

    if policy_mode and selected_file == "All":
        st.warning("Choose a specific policy document before generating an expert brief.")

    generate_brief_clicked = False
    if policy_mode:
        if st.session_state.get('policy_template_text'):
            st.caption(
                f"Template in use: {st.session_state.get('policy_template_name', 'custom template')}"
            )
        else:
            st.caption("You’ll get a Policy Vault Expert Brief: a polished, easy-to-read policy analysis")
        policy_directive = st.text_area(
            "Provide policy brief focus (optional)",
            value=st.session_state.get('policy_directive', ""),
            help="Add priorities, stakeholders, or constraints to tailor the policy brief.",
            key="policy_directive_input",
        )
        st.session_state['policy_directive'] = policy_directive
        generate_brief_clicked = st.button(
            "Generate policy brief",
            disabled=(selected_file == "All"),
            use_container_width=True,
            key="generate_policy_brief_button",
        )
    else:
        policy_directive = ""
        st.session_state.pop('policy_directive', None)
        st.session_state.pop('policy_directive_input', None)
        generate_brief_clicked = False

    user_query = None
    if not policy_mode:
        user_query = st.chat_input(placeholder="Ask anything")
    trigger_policy_request = (
        policy_mode
        and generate_brief_clicked
        and st.session_state.get('all_chunks')
        and selected_file != "All"
    )
    trigger_chat_request = user_query and st.session_state.get('all_chunks')

    if trigger_chat_request or trigger_policy_request:
        base_query = user_query or policy_directive or "Provide an expert policy analysis of the uploaded documents."
        if policy_mode and policy_directive and user_query:
            effective_query = (
                f"Primary question: {user_query}\n\nPolicy brief focus: {policy_directive}"
            )
        else:
            effective_query = base_query

        time.sleep(1)
        if selected_file != "All":
            file_name = selected_file
            relevant_chunks, _, source_files = retrieve_chunks(
                effective_query,
                st.session_state['index'],
                st.session_state['all_chunks'],
                st.session_state['all_embeddings'],
                co,
                file_name=file_name,
                all_file_names=st.session_state['all_file_names']
            )
        else:
            relevant_chunks, _, source_files = retrieve_chunks(
                effective_query,
                st.session_state['index'],
                st.session_state['all_chunks'],
                st.session_state['all_embeddings'],
                co,
                all_file_names=st.session_state.get('all_file_names')
            )

        if not relevant_chunks:
            st.warning("No content found for the selected policy. Please ensure it was processed successfully.")
        else:
            qa_context_segments = []
            policy_context_segments = []
            for chunk_text, source_name in zip(relevant_chunks, source_files):
                if source_name:
                    qa_context_segments.append(f"Policy '{source_name}': {chunk_text}")
                    policy_context_segments.append(
                        f"Source: {source_name}\nDetails: {chunk_text}"
                    )
                else:
                    qa_context_segments.append(chunk_text)
                    policy_context_segments.append(chunk_text)

            model = st.session_state.get('fine_tuned_model_id', 'command-r-08-2024')
            template_text = st.session_state.get('policy_template_text')
            used_sources = sorted({name for name in source_files if name})

            if st.session_state.get('policy_analysis_mode'):
                response_text = generate_policy_brief(
                    effective_query,
                    policy_context_segments,
                    co,
                    model=model,
                    template=template_text,
                    source_names=used_sources,
                    primary_policy=selected_file if selected_file != "All" else None,
                )
                answer = f"📄 **Policy Brief**\n\n{response_text}" if response_text else response_text
            else:
                answer = generate_answer(
                    effective_query,
                    qa_context_segments,
                    co,
                    model,
                    source_names=used_sources,
                )

            if used_sources and answer and not st.session_state.get('policy_analysis_mode'):
                answer = f"{answer}\n\n_Sources: {', '.join(used_sources)}_"

            if 'chat_history' not in st.session_state:
                st.session_state['chat_history'] = []

            if user_query and policy_mode and policy_directive:
                displayed_user_message = f"{user_query}\n\nFocus: {policy_directive}"
            else:
                displayed_user_message = user_query or (
                    f"Policy brief focus: {policy_directive}" if policy_directive else "Expert policy analysis requested."
                )

            st.session_state['chat_history'].append({"role": "user", "content": displayed_user_message})
            st.session_state['chat_history'].append({"role": "bot", "content": answer})
            st.rerun()  # Refresh to show new message

# # Footer: Evaluation Criteria & Info
# st.markdown("""
# ---
# **Precision@3**: Fraction of top-3 retrieved chunks that are relevant. High values indicate low false positives.  
# **MRR (Mean Reciprocal Rank)**: Rank of the first relevant chunk. High values mean relevant chunks are retrieved early.  

# """)
